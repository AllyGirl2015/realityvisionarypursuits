#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Criar documento
doc = Document()

# Adicionar título
title = doc.add_paragraph("The Economic Singularity")
title_run = title.runs[0]
title_run.font.size = Pt(24)
title_run.font.bold = True
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Adicionar subtítulo
subtitle = doc.add_paragraph("By Alissa M.R. Eldridge – Visionary Pursuits Series")
subtitle_run = subtitle.runs[0]
subtitle_run.font.italic = True
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Adicionar conteúdo
sections = [
    ("Introduction", [
        "The U.S. — and much of the global economy — is approaching a collapse point I call the Economic Singularity. It's the moment when wealth concentration, wage stagnation, and cost-of-living inflation converge so sharply that the majority of consumers can no longer participate in the market at all.",
        "",
        "This is not an abstract theory for some distant future. In my view, it's already underway — and the warning signs are visible in companies across nearly every sector."
    ]),
    ("Core Mechanism", [
        "The Economic Singularity follows a consistent self-destructive loop:",
        "1. Wealth concentrates at the top while real wages stagnate.",
        "2. Companies raise prices to meet profit goals, even as their core customers struggle.",
        "3. The customer base shrinks, hurting long-term revenue.",
        "4. Leaders misdiagnose this as a 'demand problem' and respond with short-term fixes — automation, layoffs, or more price hikes.",
        "5. Those fixes accelerate the collapse by further reducing the consumer base.",
        "",
        "Once this cycle reaches critical mass, the system implodes — much like matter collapsing into a black hole."
    ]),
    ("Why the Timeline Matters", [
        "I estimate a collapse window of 2–10 years from now. This is not guesswork — it's based on both historic precedent and present-day data.",
        "",
        "Historical Analogy – The Great Depression:",
        "In the 1920s, economic instability quietly grew for a decade: speculative bubbles, wage inequality, and reckless corporate practices. The 1929 stock market crash wasn't the cause of the Great Depression, but the trigger — the structural weaknesses were already there. Within three years, U.S. unemployment hit 25%.",
        "",
        "Modern Parallel:",
        "Early 2000s: Precondition phase — cost-of-living rises faster than wages, early automation adoption, growing corporate consolidation.",
        "2010s: Acceleration phase — gig economy replaces stable jobs, housing market inflation worsens, debt loads increase.",
        "2016–2024: Destabilization phase — political and policy shocks, pandemic disruption, and a rapid shift toward profit-maximizing at the expense of stability.",
        "Now: With ~15–20 years of structural degradation already behind us, the collapse phase could arrive within the next decade."
    ]),
    ("The Data That Proves the Trend", [
        "Median Household Income vs. Costs:",
        "From 2000 to 2020, median rents rose faster than incomes in 88% of U.S. counties, affecting 97% of Americans. In May 2025, rent had jumped 29% for apartments and 43% for single-family homes since 2000, while median household income grew only 23% in the same period.",
        "",
        "Real Wage Stagnation:",
        "Adjusted for inflation, U.S. median household income has been effectively flat since the late 1990s — despite GDP growth. In 2023, real median household income was $80,610, barely higher than in 2000 once inflation is accounted for.",
        "",
        "Income Inequality:",
        "By 2007, the average middle-class household income was nearly 23% lower than it could have been if growth had been evenly distributed.",
        "",
        "Conclusion from the Data: We've already burned most of the timeline between preconditions and collapse. The warning signs are no longer subtle."
    ]),
    ("Finite Money Flow", [
        "An economy's active money supply is finite at any given time. When most wealth pools at the top, circulating capital shrinks for the majority. That reduces consumer demand, which reduces revenue, which prompts companies to raise prices or cut jobs — both of which make the problem worse.",
        "",
        "It's the same logic as a star collapsing: once gravity passes a certain point, there's no stopping the implosion."
    ]),
    ("Real-World Warning Signs", [
        "McDonald's: Built for accessibility, McDonald's relies on lower- and middle-income customers. If prices ever exceed what those customers can afford — such as a single sandwich costing more than an hour's wage — many will opt for cheaper groceries.",
        "",
        "Southwest Airlines: Historically built on loyalty and simplicity, Southwest avoided many industry-standard fees. Shifting toward the same practices as competitors may yield short-term revenue gains but risks eroding the very identity that made them resilient.",
        "",
        "Boeing: Reports of cost-cutting, reduced quality control, and loss of skilled labor have already damaged Boeing's reputation.",
        "",
        "Amazon: Automation reduces payroll costs but also removes customers from the market. A robot doesn't have a paycheck to spend — meaning every replaced worker is also a lost consumer.",
        "",
        "YouTube (Google): By prioritizing advertiser comfort over creator freedom, YouTube risks hollowing out its creative base — reducing both cultural relevance and future ad revenue.",
        "",
        "Disney: An overreliance on sequels, reboots, and franchise saturation risks creative fatigue."
    ]),
    ("The Final Nail: Starving the Consumer Base", [
        "Every layoff, hiring freeze, or wage suppression in the name of 'efficiency' has two effects:",
        "1. Direct: Fewer people have money to spend.",
        "2. Compounding: Those people also spend less elsewhere, weakening other businesses in a chain reaction.",
        "",
        "By trying to save money through paying fewer people less money, companies actively shrink the very market that sustains them. Layer that onto rising prices, automation, and short-termism, and you have an economic chokehold."
    ]),
    ("Closing Thought", [
        "The Economic Singularity isn't just a theory — it's a visible trajectory backed by history, data, and ongoing corporate behavior.",
        "",
        "If the link between consumer earning power and market participation is severed, collapse isn't just likely — it's inevitable. The only question left is whether we will recognize the pattern in time to stop it."
    ])
]

for section_title, paragraphs in sections:
    heading = doc.add_heading(section_title, level=1)
    for para_text in paragraphs:
        if para_text:
            doc.add_paragraph(para_text)
        else:
            doc.add_paragraph()

# Salvar documento
doc.save("src/data/The Economic Singularity.docx")
print("Created: src/data/The Economic Singularity.docx")
